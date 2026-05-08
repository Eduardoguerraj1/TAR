from __future__ import annotations

from io import BytesIO
from typing import Any

from .model import COMPARTMENTS, INFERENTIAL_TEST_MINIMUMS
from .table_meta import tar_table_meta


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_MIME = "application/pdf"

REPORT_LEVEL_NOTE = "Report Level é o critério de notificação usado na comparação; LLD é referência de detecção e não limite de ação."
DETERMINISTIC_N_NOTE = "Nos cenários reais atuais, o n representa radionuclídeos calculados e não medições ambientais independentes."
HYPOTHETICAL_N_NOTE = "n é o número de medições sintéticas por radionuclídeo; P95 é o percentil 95 das simulações."
HYPOTHETICAL_TEST_NOTE = (
    "O teste usa o logaritmo da razão entre valor simulado e Report Level; a hipótese alternativa avalia "
    "se os resultados simulados permanecem abaixo do Report Level."
)
SENSITIVITY_NOTE = "Intervalos sintéticos demonstrativos; não substituem constantes de literatura nem conclusão regulatória."
SENSITIVITY_VARIABLES_EXPLANATION = (
    "Antes da tabela, a ideia central é simples: cada rodada do Monte Carlo sorteia um valor possível para cada "
    "variável incerta e recalcula os resultados. A distribuição triangular é usada quando há mínimo, valor mais "
    "provável e máximo, como na atividade total do TAR e na vazão de diluição. A distribuição lognormal é usada "
    "para fatores sempre positivos que variam por multiplicação, como bioacumulação e transferência para sedimento. "
    "A distribuição uniforme é usada quando só há intervalo inferior e superior, como no tempo de exposição. "
    "Como os intervalos são sintéticos, a leitura deve ser exploratória."
)
SENSITIVITY_TORNADO_EXPLANATION = (
    "O gráfico tornado é um ranking visual de influência. Para cada rodada, o sistema calcula a maior razão valor "
    "simulado / Report Level; depois mede, por correlação de Spearman, quais variáveis mais acompanham essa razão. "
    "Spearman usa a ordem dos valores e varia de -1 a +1. No tornado, usamos o valor absoluto: barras maiores "
    "indicam maior influência, seja para aumentar ou reduzir a proximidade com o Report Level."
)
SENSITIVITY_HEATMAP_EXPLANATION = (
    "O heatmap responde em quantas rodadas cada radionuclídeo e compartimento ficou acima do Report Level. "
    "A porcentagem é empírica porque vem da contagem das simulações, por exemplo 230 ultrapassagens em 10.000 "
    "rodadas equivalem a 2,3%. Células mais intensas indicam maior frequência de superação; itens sem Report "
    "Level ficam sem referência."
)
SENSITIVITY_HISTOGRAM_EXPLANATION = (
    "O histograma resume, rodada a rodada, a maior razão observada entre valor simulado e Report Level. A marca "
    "em 1,0 representa o ponto de superação da referência."
)
SENSITIVITY_INFLUENCE_EXPLANATION = (
    "A tabela de influência coloca números no que o gráfico tornado mostra visualmente. A correlação de Spearman "
    "varia de -1 a +1 e mede se, quando uma variável aumenta nas simulações, a maior razão valor simulado / Report "
    "Level também tende a aumentar ou diminuir. A ordenação usa |correlação|, por isso uma variável negativa pode "
    "aparecer como muito influente quando tem efeito redutor forte."
)
SENSITIVITY_BOXPLOT_EXPLANATION = (
    "O boxplot mostra a distribuição da maior razão simulada de forma compacta: a caixa contém os 50% centrais "
    "das rodadas, a linha interna é a mediana, as hastes mostram mínimo e máximo, e a marca de P95 destaca um "
    "cenário conservador. Ele ajuda a ver se a distribuição fica concentrada abaixo ou acima do Report Level = 1."
)
SENSITIVITY_RESULTS_EXPLANATION = (
    "A tabela de resultados consolida a faixa simulada por radionuclídeo e compartimento. O P95 é o percentil 95: "
    "após ordenar as simulações do menor para o maior, ele é o valor abaixo do qual ficam 95% dos cenários. "
    "A probabilidade empírica de ultrapassar o Report Level é a fração das simulações acima da referência, calculada "
    "como número de ultrapassagens dividido pelo total de simulações. Quando não há Report Level cadastrado, a "
    "probabilidade fica sem referência."
)
EMPIRICAL_ACTIVITY_NOTE = (
    "Valores marcados como menor que MDA são tratados como censurados: entram nas contagens de não detectados, "
    "mas não entram como valor numérico. A fração Si é calculada por amostra usando apenas radionuclídeos detectados."
)

CHART_EXPLANATIONS = {
    "Gráfico - Influência das variáveis": SENSITIVITY_TORNADO_EXPLANATION,
    "Gráfico - Probabilidade de ultrapassar Report Level": SENSITIVITY_HEATMAP_EXPLANATION,
    "Gráfico - Distribuição da maior razão simulada": SENSITIVITY_HISTOGRAM_EXPLANATION,
    "Gráfico - Boxplot da maior razão simulada": SENSITIVITY_BOXPLOT_EXPLANATION,
}


def _format_sci(value: Any) -> str:
    try:
        numeric = float(value)
    except (TypeError, ValueError):
        return "—"
    return f"{numeric:.2E}".replace("E+0", "E+").replace("E-0", "E-")


def _pdf_text(value: Any) -> str:
    from xml.sax.saxutils import escape

    return escape(str(value or ""))


def _summary_paragraph(summary: dict[str, Any]) -> str:
    scenario = summary["scenario"]
    totals = scenario["totals"]
    if scenario.get("is_hypothetical"):
        return (
            f"A avaliação TAR considera o cenário hipotético, com n = {scenario['measurement_count']} medições sintéticas "
            "de espectrometria gama por radionuclídeo da água do TAR. Cada medição alimentou uma nova simulação dos "
            "compartimentos ambientais. Os valores de concentração apresentados para comparação correspondem ao "
            f"percentil 95 das simulações. A concentração total P95 na água do mar foi "
            f"{_format_sci(totals.get('total_water_concentration_bq_m3'))} Bq/m³."
        )
    return (
        f"A avaliação TAR considera o cenário {scenario['label']}, com n = {scenario['radionuclide_count']} "
        f"radionuclídeos calculados. A concentração total na água do mar foi "
        f"{_format_sci(totals.get('total_water_concentration_bq_m3'))} Bq/m³, com atividade de "
        f"{_format_sci(totals.get('activity_bq_year'))} Bq/ano. Os valores com Report Level disponível ficaram "
        "abaixo dos níveis de notificação cadastrados."
    )


def _add_docx_table_caption(document: Any, caption: str) -> None:
    paragraph = document.add_paragraph(caption)
    for run in paragraph.runs:
        run.bold = True


def _add_docx_table_note(document: Any, note: str) -> None:
    if not note:
        return
    paragraph = document.add_paragraph(note)
    for run in paragraph.runs:
        run.italic = True


def _table_note_parts(table_key: str, *, note: str = "", source_note: str = "", unit_note: str = "") -> list[str]:
    meta = tar_table_meta(table_key)
    parts = [
        unit_note or meta.get("unit_note") or "",
        note,
        source_note or meta.get("source_note") or "",
    ]
    return [str(part) for part in parts if part]


def _add_docx_table_intro(document: Any, table_key: str) -> None:
    meta = tar_table_meta(table_key)
    document.add_paragraph(meta["lead_text"])
    if meta.get("column_notes"):
        paragraph = document.add_paragraph()
        paragraph.add_run("Leitura das colunas").bold = True
        for index, note in enumerate(meta["column_notes"], start=1):
            document.add_paragraph(f"{index}. {note}")
    _add_docx_table_caption(document, meta["display_caption"])


def _add_docx_table_notes(document: Any, table_key: str, *, note: str = "", source_note: str = "", unit_note: str = "") -> None:
    for part in _table_note_parts(table_key, note=note, source_note=source_note, unit_note=unit_note):
        _add_docx_table_note(document, part)


def _append_pdf_table_intro(story: list[Any], table_key: str, body_style: Any, caption_style: Any) -> None:
    from reportlab.platypus import Paragraph

    meta = tar_table_meta(table_key)
    story.append(Paragraph(_pdf_text(meta["lead_text"]), body_style))
    if meta.get("column_notes"):
        story.append(Paragraph("<b>Leitura das colunas</b>", body_style))
        for index, note in enumerate(meta["column_notes"], start=1):
            story.append(Paragraph(f"{index}. {_pdf_text(note)}", body_style))
    story.append(Paragraph(_pdf_text(meta["display_caption"]), caption_style))


def _append_pdf_table_notes(
    story: list[Any],
    table_key: str,
    note_style: Any,
    *,
    note: str = "",
    source_note: str = "",
    unit_note: str = "",
) -> None:
    from reportlab.platypus import Paragraph

    for part in _table_note_parts(table_key, note=note, source_note=source_note, unit_note=unit_note):
        story.append(Paragraph(_pdf_text(part), note_style))


def _load_pillow_font(size: int, *, bold: bool = False) -> Any:
    from PIL import ImageFont

    candidates = (
        ["C:/Windows/Fonts/arialbd.ttf", "DejaVuSans-Bold.ttf"]
        if bold
        else ["C:/Windows/Fonts/arial.ttf", "DejaVuSans.ttf"]
    )
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except Exception:
            continue
    return ImageFont.load_default()


def _chart_png_buffer(width: int, height: int) -> tuple[Any, Any, Any, BytesIO]:
    from PIL import Image, ImageDraw

    image = Image.new("RGB", (width, height), "#ffffff")
    draw = ImageDraw.Draw(image)
    buffer = BytesIO()
    return image, draw, buffer, buffer


def _finish_png(image: Any, buffer: BytesIO) -> BytesIO:
    image.save(buffer, format="PNG")
    buffer.seek(0)
    return buffer


def _sensitivity_tornado_png(sensitivity: dict[str, Any]) -> tuple[BytesIO, int, int]:
    rows = (sensitivity.get("chart_payloads") or {}).get("tornado", {}).get("rows") or []
    width = 1050
    left = 360
    right = 120
    top = 56
    row_height = 42
    height = top + max(1, len(rows)) * row_height + 34
    image, draw, buffer, _ = _chart_png_buffer(width, height)
    font = _load_pillow_font(18)
    small = _load_pillow_font(15)
    title_font = _load_pillow_font(22, bold=True)
    draw.text((24, 18), "Influência das variáveis", fill="#17313a", font=title_font)
    bar_width = width - left - right
    draw.line((left, 44, left + bar_width, 44), fill="#c9d7dc", width=2)
    for index, row in enumerate(rows):
        y = top + index * row_height
        absolute = max(0.0, min(1.0, float(row.get("absolute_correlation") or 0.0)))
        corr = float(row.get("correlation") or 0.0)
        color = "#27667b" if corr >= 0 else "#b8672a"
        bar = max(4, int(absolute * bar_width))
        draw.text((24, y + 8), str(row.get("label") or ""), fill="#17313a", font=font)
        draw.rounded_rectangle((left, y + 7, left + bar, y + 29), radius=6, fill=color)
        draw.text((left + bar + 10, y + 8), str(row.get("correlation_text") or "—"), fill="#516873", font=small)
    return _finish_png(image, buffer), width, height


def _heatmap_color(probability: Any) -> str:
    if probability is None:
        return "#eef2f4"
    value = max(0.0, min(1.0, float(probability)))
    if value <= 0.01:
        return "#e8f3ee"
    if value <= 0.05:
        return "#c7e6d2"
    if value <= 0.20:
        return "#f2d27a"
    if value <= 0.50:
        return "#df8a4a"
    return "#b8423a"


def _sensitivity_heatmap_png(sensitivity: dict[str, Any]) -> tuple[BytesIO, int, int]:
    payload = (sensitivity.get("chart_payloads") or {}).get("heatmap") or {}
    radionuclides = payload.get("radionuclides") or []
    compartments = payload.get("compartments") or []
    cells = {
        (cell.get("radionuclide"), cell.get("compartment_key")): cell
        for cell in payload.get("cells") or []
    }
    cell_width = 150
    cell_height = 36
    left = 130
    top = 70
    width = left + len(compartments) * cell_width + 40
    height = top + len(radionuclides) * cell_height + 36
    image, draw, buffer, _ = _chart_png_buffer(width, height)
    font = _load_pillow_font(16)
    small = _load_pillow_font(14)
    title_font = _load_pillow_font(22, bold=True)
    draw.text((24, 18), "Probabilidade de ultrapassar Report Level", fill="#17313a", font=title_font)
    for col, compartment in enumerate(compartments):
        x = left + col * cell_width + 10
        draw.text((x, 48), str(compartment.get("label") or ""), fill="#17313a", font=small)
    for row_index, radionuclide in enumerate(radionuclides):
        y = top + row_index * cell_height
        draw.text((24, y + 8), str(radionuclide), fill="#17313a", font=small)
        for col, compartment in enumerate(compartments):
            x = left + col * cell_width
            cell = cells.get((radionuclide, compartment.get("key"))) or {}
            draw.rounded_rectangle((x, y, x + cell_width - 8, y + cell_height - 6), radius=6, fill=_heatmap_color(cell.get("probability")))
            draw.text((x + 42, y + 8), str(cell.get("probability_text") or "—"), fill="#17313a", font=font)
    return _finish_png(image, buffer), width, height


def _sensitivity_histogram_png(sensitivity: dict[str, Any]) -> tuple[BytesIO, int, int]:
    payload = (sensitivity.get("chart_payloads") or {}).get("histogram") or {}
    bins = payload.get("bins") or []
    width = 1050
    height = 360
    left = 72
    right = 36
    top = 68
    bottom = 72
    plot_width = width - left - right
    plot_height = height - top - bottom
    image, draw, buffer, _ = _chart_png_buffer(width, height)
    font = _load_pillow_font(15)
    title_font = _load_pillow_font(22, bold=True)
    draw.text((24, 18), "Distribuição da maior razão simulada", fill="#17313a", font=title_font)
    draw.rounded_rectangle((left, top, left + plot_width, top + plot_height), radius=8, outline="#d8e2e6", fill="#f7fafb")
    max_count = max([int(item.get("count") or 0) for item in bins] + [1])
    lower = min([float(item.get("start") or 0) for item in bins] + [0.0])
    upper = max([float(item.get("end") or 1) for item in bins] + [1.0])
    if upper <= lower:
        upper = lower + 1.0
    bar_slot = plot_width / max(1, len(bins))
    for index, item in enumerate(bins):
        count = int(item.get("count") or 0)
        bar_height = 0 if max_count <= 0 else (count / max_count) * (plot_height - 16)
        x = left + index * bar_slot
        y = top + plot_height - bar_height
        draw.rectangle((x + 3, y, x + bar_slot - 4, top + plot_height), fill="#27667b")
        if index % 3 == 0:
            draw.text((x + 2, height - 48), str(item.get("label") or ""), fill="#607782", font=font)
    reference = float(payload.get("reference") or 1.0)
    if lower <= reference <= upper:
        ref_x = left + ((reference - lower) / (upper - lower)) * plot_width
        draw.line((ref_x, top, ref_x, top + plot_height), fill="#b8423a", width=3)
        draw.text((ref_x + 8, top + 8), "Report Level", fill="#b8423a", font=font)
    draw.text((left + 180, height - 22), "Maior razão valor simulado / Report Level por simulação", fill="#17313a", font=font)
    return _finish_png(image, buffer), width, height


def _sensitivity_boxplot_png(sensitivity: dict[str, Any]) -> tuple[BytesIO, int, int]:
    payload = (sensitivity.get("chart_payloads") or {}).get("boxplot") or {}
    stats = payload.get("stats") or {}
    values = [
        stats.get("min"),
        stats.get("q1"),
        stats.get("median"),
        stats.get("q3"),
        stats.get("max"),
        stats.get("p95"),
        payload.get("reference") or 1.0,
    ]
    numeric = [float(value) for value in values if value is not None]
    width = 1050
    height = 260
    image, draw, buffer, _ = _chart_png_buffer(width, height)
    font = _load_pillow_font(15)
    title_font = _load_pillow_font(22, bold=True)
    draw.text((24, 18), "Boxplot da maior razão simulada", fill="#17313a", font=title_font)
    if not numeric:
        return _finish_png(image, buffer), width, height
    left = 90
    right = 48
    axis_y = 128
    plot_width = width - left - right
    lower = min(0.0, min(numeric))
    upper = max(numeric)
    if upper <= lower:
        upper = lower + 1.0

    def x_at(value: Any) -> float:
        return left + ((float(value) - lower) / (upper - lower)) * plot_width

    min_x = x_at(stats.get("min") or lower)
    q1_x = x_at(stats.get("q1") or lower)
    median_x = x_at(stats.get("median") or lower)
    q3_x = x_at(stats.get("q3") or lower)
    max_x = x_at(stats.get("max") or upper)
    p95_x = x_at(stats.get("p95") or upper)
    ref_x = x_at(payload.get("reference") or 1.0)
    draw.line((left, axis_y, left + plot_width, axis_y), fill="#d8e2e6", width=2)
    draw.line((min_x, axis_y, q1_x, axis_y), fill="#27667b", width=4)
    draw.line((q3_x, axis_y, max_x, axis_y), fill="#27667b", width=4)
    draw.line((min_x, axis_y - 28, min_x, axis_y + 28), fill="#27667b", width=4)
    draw.line((max_x, axis_y - 28, max_x, axis_y + 28), fill="#27667b", width=4)
    draw.rounded_rectangle((q1_x, axis_y - 34, max(q1_x + 3, q3_x), axis_y + 34), radius=8, fill="#d9ecf1", outline="#27667b", width=4)
    draw.line((median_x, axis_y - 34, median_x, axis_y + 34), fill="#17313a", width=5)
    draw.line((p95_x, axis_y - 42, p95_x, axis_y + 42), fill="#6b4aa0", width=3)
    draw.text((p95_x + 8, axis_y - 62), f"P95 {stats.get('p95_text') or '—'}", fill="#6b4aa0", font=font)
    draw.line((ref_x, 58, ref_x, height - 46), fill="#b8423a", width=3)
    draw.text((ref_x + 8, 54), "Report Level = 1", fill="#b8423a", font=font)
    draw.text((min_x - 20, height - 46), f"mín {stats.get('min_text') or '—'}", fill="#607782", font=font)
    draw.text((median_x - 50, height - 76), f"mediana {stats.get('median_text') or '—'}", fill="#17313a", font=font)
    draw.text((max_x - 40, height - 46), f"máx {stats.get('max_text') or '—'}", fill="#607782", font=font)
    draw.text((left + 290, height - 18), "Maior razão valor simulado / Report Level", fill="#17313a", font=font)
    return _finish_png(image, buffer), width, height


def _sensitivity_chart_images(sensitivity: dict[str, Any]) -> list[tuple[str, BytesIO, int, int]]:
    return [
        ("Gráfico - Influência das variáveis", *_sensitivity_tornado_png(sensitivity)),
        ("Gráfico - Probabilidade de ultrapassar Report Level", *_sensitivity_heatmap_png(sensitivity)),
        ("Gráfico - Distribuição da maior razão simulada", *_sensitivity_histogram_png(sensitivity)),
        ("Gráfico - Boxplot da maior razão simulada", *_sensitivity_boxplot_png(sensitivity)),
    ]


def _add_docx_empirical_activity_section(document: Any, empirical: dict[str, Any]) -> None:
    if not empirical:
        return

    document.add_heading("Dados reais de atividade total TAR", level=1)
    document.add_paragraph(empirical.get("narrative_text") or "")
    document.add_paragraph(" | ".join(f"{card.get('label')}: {card.get('value')}" for card in empirical.get("cards") or []))

    _add_docx_table_intro(document, "empirical_groups")
    group_table = document.add_table(rows=1, cols=8)
    group_table.style = "Table Grid"
    for index, header in enumerate(["Grupo", "Amostras", "A-TAR detectado", "A-TAR < MDA", "A-TAR ausente", "Média", "Mediana", "P95"]):
        group_table.rows[0].cells[index].text = header
    for item in empirical.get("groups") or []:
        cells = group_table.add_row().cells
        cells[0].text = item.get("group", "")
        cells[1].text = str(item.get("sample_count") or 0)
        cells[2].text = str(item.get("activity_detected_count") or 0)
        cells[3].text = str(item.get("activity_censored_count") or 0)
        cells[4].text = str(item.get("activity_missing_count") or 0)
        cells[5].text = item.get("activity_mean_text", "—")
        cells[6].text = item.get("activity_median_text", "—")
        cells[7].text = item.get("activity_p95_text", "—")
    _add_docx_table_notes(document, "empirical_groups", note=empirical.get("mda_policy") or EMPIRICAL_ACTIVITY_NOTE)

    _add_docx_table_intro(document, "empirical_radionuclides")
    radionuclide_table = document.add_table(rows=1, cols=11)
    radionuclide_table.style = "Table Grid"
    for index, header in enumerate(["Grupo", "Radionuclídeo", "Status no modelo", "Amostras", "Detectados", "< MDA", "Ausentes", "Taxa detectada", "Média", "Mediana", "P95"]):
        radionuclide_table.rows[0].cells[index].text = header
    for item in empirical.get("radionuclide_rows") or []:
        cells = radionuclide_table.add_row().cells
        cells[0].text = item.get("group", "")
        cells[1].text = item.get("radionuclide", "")
        cells[2].text = item.get("model_status", "")
        cells[3].text = str(item.get("sample_count") or 0)
        cells[4].text = str(item.get("detected_count") or 0)
        cells[5].text = str(item.get("censored_count") or 0)
        cells[6].text = str(item.get("missing_count") or 0)
        cells[7].text = item.get("detected_rate_text", "—")
        cells[8].text = item.get("mean_text", "—")
        cells[9].text = item.get("median_text", "—")
        cells[10].text = item.get("p95_text", "—")
    _add_docx_table_notes(
        document,
        "empirical_radionuclides",
        note="Nb-95 permanece como observado não modelado porque não há linha equivalente na fórmula da planilha TAR atual.",
    )

    _add_docx_table_intro(document, "empirical_modeled")
    modeled_table = document.add_table(rows=1, cols=11)
    modeled_table.style = "Table Grid"
    for index, header in enumerate(["Grupo", "Radionuclídeo", "Compartimento", "n", "Média", "Mediana", "P95", "Report Level", "P95/RL", "Status", "Freq. > RL"]):
        modeled_table.rows[0].cells[index].text = header
    for item in empirical.get("modeled_compartment_rows") or []:
        cells = modeled_table.add_row().cells
        cells[0].text = item.get("group", "")
        cells[1].text = item.get("radionuclide", "")
        cells[2].text = item.get("compartment", "")
        cells[3].text = str(item.get("n") or "—")
        cells[4].text = item.get("mean_text", "—")
        cells[5].text = item.get("median_text", "—")
        cells[6].text = item.get("p95_text", "—")
        cells[7].text = item.get("report_level_text", "—")
        cells[8].text = item.get("report_level_p95_ratio_text", "—")
        cells[9].text = item.get("report_level_status", "—")
        cells[10].text = item.get("report_level_exceedance_rate_text", "—")
    _add_docx_table_notes(
        document,
        "empirical_modeled",
        note="As frações Si foram calculadas por amostra a partir dos radionuclídeos detectados; < MDA> não participa do denominador.",
    )

    _add_docx_table_intro(document, "empirical_inferential")
    inferential_table = document.add_table(rows=1, cols=12)
    inferential_table.style = "Table Grid"
    headers = [
        "Grupo",
        "Radionuclídeo",
        "Compartimento",
        "n",
        "Report Level",
        "P95 razão",
        "Ultrapassagens",
        "Freq. > RL",
        "IC95% freq.",
        "Teste",
        "p-value",
        "Conclusão",
    ]
    for index, header in enumerate(headers):
        inferential_table.rows[0].cells[index].text = header
    for item in empirical.get("inferential_rows") or []:
        cells = inferential_table.add_row().cells
        cells[0].text = item.get("group", "")
        cells[1].text = item.get("radionuclide", "")
        cells[2].text = item.get("compartment", "")
        cells[3].text = str(item.get("n") or "—")
        cells[4].text = item.get("reference_value_text", "—")
        cells[5].text = item.get("p95_ratio_text", "—")
        cells[6].text = str(item.get("exceedance_count") if item.get("exceedance_count") is not None else "—")
        cells[7].text = item.get("exceedance_rate_text", "—")
        cells[8].text = item.get("exceedance_ci95_text", "—")
        cells[9].text = item.get("test_label", "—")
        cells[10].text = item.get("p_value_text", "—")
        cells[11].text = item.get("conclusion") or item.get("reason", "")
    _add_docx_table_notes(
        document,
        "empirical_inferential",
        note="O Report Level é referência fixa; a incerteza estatística é estimada a partir das amostras reais calculadas pela fórmula.",
    )


def _add_docx_sensitivity_section(document: Any, sensitivity: dict[str, Any]) -> None:
    if not sensitivity:
        return
    from docx.shared import Inches

    document.add_heading("Análise de sensibilidade Monte Carlo", level=1)
    document.add_paragraph(sensitivity.get("narrative_text") or "")
    document.add_paragraph(" | ".join(f"{card.get('label')}: {card.get('value')}" for card in sensitivity.get("cards") or []))

    _add_docx_table_intro(document, "sensitivity_variables")
    variable_table = document.add_table(rows=1, cols=6)
    variable_table.style = "Table Grid"
    for index, header in enumerate(["Variável", "Distribuição", "Parâmetros", "Base", "Unidade", "Uso no modelo"]):
        variable_table.rows[0].cells[index].text = header
    for variable in sensitivity.get("variables") or []:
        cells = variable_table.add_row().cells
        cells[0].text = variable.get("label", "")
        cells[1].text = variable.get("distribution", "")
        cells[2].text = variable.get("parameters", "")
        cells[3].text = variable.get("base_value_text", "—")
        cells[4].text = variable.get("unit", "")
        cells[5].text = variable.get("description", "")
    _add_docx_table_notes(document, "sensitivity_variables", note=sensitivity.get("source_note") or SENSITIVITY_NOTE)

    for title, image_buffer, _width, _height in _sensitivity_chart_images(sensitivity):
        document.add_paragraph(CHART_EXPLANATIONS.get(title, ""))
        _add_docx_table_caption(document, title)
        document.add_picture(image_buffer, width=Inches(6.6))

    _add_docx_table_intro(document, "sensitivity_influence")
    influence_table = document.add_table(rows=1, cols=4)
    influence_table.style = "Table Grid"
    for index, header in enumerate(["Variável", "Spearman", "|Correlação|", "Sentido"]):
        influence_table.rows[0].cells[index].text = header
    for item in sensitivity.get("influence_rows") or []:
        cells = influence_table.add_row().cells
        cells[0].text = item.get("label", "")
        cells[1].text = item.get("correlation_text", "—")
        cells[2].text = item.get("absolute_correlation_text", "—")
        cells[3].text = item.get("direction", "—")
    _add_docx_table_notes(
        document,
        "sensitivity_influence",
        note=(
            "Correlação positiva indica que valores maiores da variável tendem a aumentar a maior razão contra o Report Level; "
            "correlação negativa indica que valores maiores tendem a reduzir essa razão."
        ),
    )

    _add_docx_table_intro(document, "sensitivity_results")
    result_table = document.add_table(rows=1, cols=9)
    result_table.style = "Table Grid"
    for index, header in enumerate(["Radionuclídeo", "Compartimento", "Média", "Mínimo", "Máximo", "P95", "Report Level", "P95/RL", "Prob. > RL"]):
        result_table.rows[0].cells[index].text = header
    for item in sensitivity.get("summary_rows") or []:
        cells = result_table.add_row().cells
        cells[0].text = item.get("radionuclide", "")
        cells[1].text = item.get("compartment", "")
        cells[2].text = item.get("mean_text", "—")
        cells[3].text = item.get("min_text", "—")
        cells[4].text = item.get("max_text", "—")
        cells[5].text = item.get("p95_text", "—")
        cells[6].text = item.get("report_level_text", "—")
        cells[7].text = item.get("p95_report_level_ratio_text", "—")
        cells[8].text = item.get("exceedance_probability_text", "—")
    _add_docx_table_notes(
        document,
        "sensitivity_results",
        note=(
            "A probabilidade empírica é calculada como número de simulações acima do Report Level dividido pelo número total de simulações. "
            "Quando não há Report Level cadastrado, a probabilidade fica sem referência."
        ),
    )


def _add_docx_stat_source_table(document: Any, statistical: dict[str, Any], rows_key: str, table_key: str) -> None:
    _add_docx_table_intro(document, table_key)
    if rows_key == "norm_rows":
        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = ["Radionuclídeo", "Compartimento", "Referência", "Valor normativo", "Unidade"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for item in statistical.get(rows_key) or []:
            cells = table.add_row().cells
            cells[0].text = item.get("radionuclide", "")
            cells[1].text = item.get("compartment", "")
            cells[2].text = item.get("reference", "")
            cells[3].text = item.get("value_text", "—")
            cells[4].text = item.get("unit", "")
    else:
        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = ["Radionuclídeo", "Compartimento", "Valor base", "Unidade", "Origem"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for item in statistical.get(rows_key) or []:
            cells = table.add_row().cells
            cells[0].text = item.get("radionuclide", "")
            cells[1].text = item.get("compartment", "")
            cells[2].text = item.get("value_text", "—")
            cells[3].text = item.get("unit", "")
            cells[4].text = item.get("method", "")
    _add_docx_table_notes(
        document,
        table_key,
        note="Os valores normativos são fixos; os valores calculados e ERICA são bases determinísticas ou estimadas, sem replicações aleatórias.",
    )


def _add_docx_statistical_section(document: Any, statistical: dict[str, Any]) -> None:
    if not statistical:
        return

    document.add_heading("Estatística calculado vs ERICA vs norma", level=1)
    document.add_paragraph(statistical.get("narrative_text") or "")
    _add_docx_table_note(
        document,
        "Esta seção não usa replicações aleatórias. A norma é referência fixa; os valores do ERICA Tool são estimativas para visualização até substituição por saídas reais.",
    )

    _add_docx_stat_source_table(document, statistical, "calculated_rows", "stat_calculated")
    _add_docx_stat_source_table(document, statistical, "erica_rows", "stat_erica")
    _add_docx_stat_source_table(document, statistical, "norm_rows", "stat_norms")

    _add_docx_table_intro(document, "stat_descriptive")
    descriptive = document.add_table(rows=1, cols=11)
    descriptive.style = "Table Grid"
    for index, header in enumerate(["Conjunto", "Radionuclídeo", "Compartimento", "n", "Média", "Mediana", "Desvio-padrão", "Q1", "Q3", "P95", "CV"]):
        descriptive.rows[0].cells[index].text = header
    for item in statistical.get("descriptive_rows") or []:
        cells = descriptive.add_row().cells
        cells[0].text = item.get("dataset", "")
        cells[1].text = item.get("radionuclide", "")
        cells[2].text = item.get("compartment", "")
        cells[3].text = str(item.get("n") or "—")
        cells[4].text = item.get("mean_text", "—")
        cells[5].text = item.get("median_text", "—")
        cells[6].text = item.get("stdev_text", "—")
        cells[7].text = item.get("q1_text", "—")
        cells[8].text = item.get("q3_text", "—")
        cells[9].text = item.get("p95_text", "—")
        cells[10].text = item.get("cv_text", "—")
    _add_docx_table_notes(document, "stat_descriptive", note="A estatística usa os valores disponíveis por radionuclídeo e compartimento; não há aumento artificial de N por sorteio.")

    _add_docx_table_intro(document, "stat_paired")
    paired = document.add_table(rows=1, cols=8)
    paired.style = "Table Grid"
    for index, header in enumerate(["Escopo", "Compartimento", "n", "Mediana calculado/ERICA", "IC95% da razão média", "Teste", "p-value", "Conclusão"]):
        paired.rows[0].cells[index].text = header
    for item in statistical.get("paired_comparison_rows") or []:
        cells = paired.add_row().cells
        cells[0].text = item.get("scope") or item.get("radionuclide", "")
        cells[1].text = item.get("compartment", "")
        cells[2].text = str(item.get("n") or "—")
        cells[3].text = item.get("median_ratio_text", "—")
        cells[4].text = f"{item.get('ci95_low_text', '—')} - {item.get('ci95_high_text', '—')}"
        cells[5].text = item.get("test_label", "—")
        cells[6].text = item.get("p_value_text", "—")
        cells[7].text = item.get("conclusion") or item.get("reason", "")
    _add_docx_table_notes(document, "stat_paired", note="A comparação pareada usa log(calculado / ERICA) e tem finalidade exploratória; os valores do ERICA são estimativas visuais.")


def build_tar_docx_payload(summary: dict[str, Any]) -> tuple[bytes, str, str]:
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception as exc:  # pragma: no cover - exercised only when dependency is missing
        raise RuntimeError("A biblioteca python-docx não está disponível para gerar o DOCX TAR.") from exc

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    scenario = summary["scenario"]
    document.add_heading("Relatório TAR", level=0)
    document.add_paragraph(_summary_paragraph(summary))
    document.add_paragraph(
        (
            "A inferência do relatório usa as amostras reais do TAR - Afluente calculadas pela fórmula da planilha. "
            "Report Level e LLD permanecem referências fixas; p-values, quando exibidos, descrevem os resultados por amostra, não a norma."
        )
        if not scenario.get("is_hypothetical")
        else "O n do cenário hipotético corresponde ao número de medições sintéticas por radionuclídeo. O teste inferencial foi aplicado sobre o logaritmo da razão simulado/Report Level."
    )

    document.add_heading("Referências disponíveis", level=1)
    _add_docx_table_intro(document, "reference_counts")
    reference_counts = document.add_table(rows=1, cols=4)
    reference_counts.style = "Table Grid"
    for index, header in enumerate(["Compartimento", "Report Level", "LLD", "Radionuclídeos com referência"]):
        reference_counts.rows[0].cells[index].text = header
    for compartment in COMPARTMENTS:
        count = scenario["reference_counts"][compartment["key"]]
        cells = reference_counts.add_row().cells
        cells[0].text = count["label"]
        cells[1].text = str(count["report_level"])
        cells[2].text = str(count["lld"])
        cells[3].text = ", ".join(count["report_level_radionuclides"]) or "sem referência"
    _add_docx_table_notes(document, "reference_counts", note=REPORT_LEVEL_NOTE)

    document.add_heading("Concentrações calculadas", level=1)
    _add_docx_table_intro(document, "concentrations")
    table = document.add_table(rows=1, cols=1 + len(COMPARTMENTS))
    table.style = "Table Grid"
    headers = ["Radionuclídeo", *[f"{comp['label']} ({comp['unit']})" for comp in COMPARTMENTS]]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for item in scenario["rows"]:
        cells = table.add_row().cells
        cells[0].text = item["radionuclide"]
        for index, compartment in enumerate(COMPARTMENTS, start=1):
            cells[index].text = item["compartments"][compartment["key"]]["value_text"]
    _add_docx_table_notes(document, "concentrations", note=HYPOTHETICAL_N_NOTE if scenario.get("is_hypothetical") else DETERMINISTIC_N_NOTE)

    document.add_heading("Comparação com Report Level e LLD", level=1)
    _add_docx_table_intro(document, "reference_results")
    reference_table = document.add_table(rows=1, cols=9)
    reference_table.style = "Table Grid"
    reference_headers = [
        "Radionuclídeo",
        "Compartimento",
        "Concentração",
        "Report Level",
        "Razão valor/Report Level",
        "Status Report Level",
        "LLD",
        "Razão valor/LLD",
        "Status LLD",
    ]
    for index, header in enumerate(reference_headers):
        reference_table.rows[0].cells[index].text = header
    for item in scenario["rows"]:
        for compartment in COMPARTMENTS:
            data = item["compartments"][compartment["key"]]
            cells = reference_table.add_row().cells
            cells[0].text = item["radionuclide"]
            cells[1].text = compartment["label"]
            cells[2].text = data["value_text"]
            cells[3].text = data["report_level_text"]
            cells[4].text = data["report_level_ratio_text"]
            cells[5].text = data["report_level_status"]
            cells[6].text = data["lld_text"]
            cells[7].text = data["lld_ratio_text"]
            cells[8].text = data["lld_status"]
    _add_docx_table_notes(document, "reference_results", note=REPORT_LEVEL_NOTE)

    if scenario.get("is_hypothetical"):
        document.add_heading("Medições sintéticas da água do TAR", level=1)
        _add_docx_table_intro(document, "hypothetical_measurements")
        measurement_table = document.add_table(rows=1, cols=6)
        measurement_table.style = "Table Grid"
        for index, header in enumerate(["Radionuclídeo", "n", "Base TAR", "Média", "Mediana", "P95"]):
            measurement_table.rows[0].cells[index].text = header
        for item in scenario["rows"]:
            cells = measurement_table.add_row().cells
            measurement = item["measurement_summary"]
            cells[0].text = item["radionuclide"]
            cells[1].text = str(item["measurement_count"])
            cells[2].text = item["source_concentration_bq_m3_text"]
            cells[3].text = measurement["mean_text"]
            cells[4].text = measurement["median_text"]
            cells[5].text = measurement["p95_text"]
        _add_docx_table_notes(document, "hypothetical_measurements", note=HYPOTHETICAL_N_NOTE)

        document.add_heading("Teste estatístico do cenário hipotético", level=1)
        document.add_paragraph(scenario.get("statistical_text") or "")
        _add_docx_table_intro(document, "hypothetical_tests")
        test_table = document.add_table(rows=1, cols=9)
        test_table.style = "Table Grid"
        for index, header in enumerate(["Radionuclídeo", "Compartimento", "n sim.", "Shapiro-Wilk", "Teste usado", "p-value", "P95 simulado / Report Level", "Acima do limite", "Conclusão"]):
            test_table.rows[0].cells[index].text = header
        for result in scenario.get("statistical_tests", []):
            cells = test_table.add_row().cells
            cells[0].text = result.get("radionuclide", "")
            cells[1].text = result.get("compartment", "")
            cells[2].text = str(result.get("n") or "—")
            cells[3].text = result.get("shapiro_p_text") or "—"
            cells[4].text = result.get("test_label") or "—"
            cells[5].text = result.get("p_value_text") or "—"
            cells[6].text = result.get("p95_ratio_text") or "—"
            cells[7].text = str(result.get("exceedance_count") or 0)
            cells[8].text = result.get("conclusion") or ""
        _add_docx_table_notes(document, "hypothetical_tests", note=HYPOTHETICAL_TEST_NOTE)

    _add_docx_empirical_activity_section(document, scenario.get("empirical_activity_statistics") or {})

    _add_docx_statistical_section(document, scenario.get("statistical_comparison") or {})

    _add_docx_sensitivity_section(document, scenario.get("sensitivity") or {})

    document.add_heading("Suficiência estatística", level=1)
    document.add_paragraph(summary["inferential_assessment"]["reason"])
    _add_docx_table_intro(document, "minimums")
    minimums = document.add_table(rows=1, cols=3)
    minimums.style = "Table Grid"
    for index, header in enumerate(["Teste", "Mínimo técnico", "Recomendado para relatório"]):
        minimums.rows[0].cells[index].text = header
    for item in INFERENTIAL_TEST_MINIMUMS:
        cells = minimums.add_row().cells
        cells[0].text = item["test"]
        cells[1].text = item["technical_minimum"]
        cells[2].text = item["recommended"]
    _add_docx_table_notes(document, "minimums", note=DETERMINISTIC_N_NOTE)

    output = BytesIO()
    document.save(output)
    filename = f"relatorio_tar_{summary['selected_scenario']}.docx"
    return output.getvalue(), filename, DOCX_MIME


def build_tar_pdf_payload(summary: dict[str, Any]) -> tuple[bytes, str, str]:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import Image as RLImage
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    except Exception as exc:  # pragma: no cover - exercised only when dependency is missing
        raise RuntimeError("A biblioteca reportlab não está disponível para gerar o PDF TAR.") from exc

    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        leftMargin=1.6 * cm,
        rightMargin=1.6 * cm,
        topMargin=1.4 * cm,
        bottomMargin=1.4 * cm,
    )
    styles = getSampleStyleSheet()
    body_style = ParagraphStyle("TarBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=9, leading=12)
    note_style = ParagraphStyle("TarNote", parent=body_style, fontSize=8, textColor=colors.HexColor("#607782"), leading=10)
    heading_style = ParagraphStyle("TarHeading", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=15, spaceAfter=8)
    subheading_style = ParagraphStyle("TarSubheading", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=12, spaceBefore=12, spaceAfter=6)
    caption_style = ParagraphStyle("TarTableCaption", parent=body_style, fontName="Helvetica-Bold", textColor=colors.HexColor("#203f49"), spaceBefore=4, spaceAfter=4)
    story: list[Any] = [
        Paragraph("Relatório TAR", heading_style),
        Paragraph(_summary_paragraph(summary), body_style),
        Spacer(1, 8),
    ]

    def apply_table_style(table: Any, *, font_size: float = 7) -> Any:
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eaf1f3")),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#c9d7dc")),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), font_size),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        return table

    scenario = summary["scenario"]

    story.append(Paragraph("Referências disponíveis", subheading_style))
    _append_pdf_table_intro(story, "reference_counts", body_style, caption_style)
    reference_count_rows = [["Compartimento", "Report Level", "LLD", "Radionuclídeos com referência"]]
    for compartment in COMPARTMENTS:
        count = scenario["reference_counts"][compartment["key"]]
        reference_count_rows.append(
            [
                count["label"],
                str(count["report_level"]),
                str(count["lld"]),
                ", ".join(count["report_level_radionuclides"]) or "sem referência",
            ]
        )
    reference_count_table = apply_table_style(
        Table(reference_count_rows, colWidths=[3.0 * cm, 2.4 * cm, 2.0 * cm, 8.6 * cm], repeatRows=1),
        font_size=7,
    )
    story.append(reference_count_table)
    _append_pdf_table_notes(story, "reference_counts", note_style, note=REPORT_LEVEL_NOTE)
    story.append(Spacer(1, 8))

    story.append(Paragraph("Concentrações calculadas", subheading_style))
    _append_pdf_table_intro(story, "concentrations", body_style, caption_style)
    concentration_rows = [["Radionuclídeo", *[f"{comp['label']} ({comp['unit']})" for comp in COMPARTMENTS]]]
    for item in scenario["rows"]:
        concentration_rows.append(
            [
                item["radionuclide"],
                *[item["compartments"][compartment["key"]]["value_text"] for compartment in COMPARTMENTS],
            ]
        )
    concentration_table = apply_table_style(
        Table(concentration_rows, colWidths=[2.2 * cm, 3.3 * cm, 3.3 * cm, 3.4 * cm, 3.3 * cm], repeatRows=1),
        font_size=6.3,
    )
    story.append(concentration_table)
    _append_pdf_table_notes(
        story,
        "concentrations",
        note_style,
        note=HYPOTHETICAL_N_NOTE if scenario.get("is_hypothetical") else DETERMINISTIC_N_NOTE,
    )
    story.append(Spacer(1, 8))

    story.append(Paragraph("Comparação com Report Level e LLD", subheading_style))
    _append_pdf_table_intro(story, "reference_results", body_style, caption_style)
    reference_rows = [["Radionuclídeo", "Compart.", "Concentração", "Report Level", "Razão RL", "Status RL", "LLD", "Razão LLD", "Status LLD"]]
    for item in scenario["rows"]:
        for compartment in COMPARTMENTS:
            data = item["compartments"][compartment["key"]]
            reference_rows.append(
                [
                    item["radionuclide"],
                    compartment["label"],
                    data["value_text"],
                    data["report_level_text"],
                    data["report_level_ratio_text"],
                    data["report_level_status"],
                    data["lld_text"],
                    data["lld_ratio_text"],
                    data["lld_status"],
                ]
            )
    reference_table = apply_table_style(
        Table(
            reference_rows,
            colWidths=[1.6 * cm, 1.7 * cm, 2.0 * cm, 2.0 * cm, 1.5 * cm, 1.5 * cm, 1.8 * cm, 1.5 * cm, 1.5 * cm],
            repeatRows=1,
        ),
        font_size=4.8,
    )
    story.append(reference_table)
    _append_pdf_table_notes(story, "reference_results", note_style, note=REPORT_LEVEL_NOTE)
    story.append(Spacer(1, 8))

    if scenario.get("is_hypothetical"):
        story.append(Paragraph("Medições sintéticas da água do TAR", subheading_style))
        _append_pdf_table_intro(story, "hypothetical_measurements", body_style, caption_style)
        measurement_rows = [["Radionuclídeo", "n", "Base TAR", "Média", "Mediana", "P95"]]
        for item in scenario["rows"]:
            measurement = item["measurement_summary"]
            measurement_rows.append(
                [
                    item["radionuclide"],
                    str(item["measurement_count"]),
                    item["source_concentration_bq_m3_text"],
                    measurement["mean_text"],
                    measurement["median_text"],
                    measurement["p95_text"],
                ]
            )
        measurement_table = apply_table_style(
            Table(measurement_rows, colWidths=[2.3 * cm, 1.0 * cm, 3.1 * cm, 3.1 * cm, 3.1 * cm, 3.1 * cm], repeatRows=1),
            font_size=6.2,
        )
        story.append(measurement_table)
        _append_pdf_table_notes(story, "hypothetical_measurements", note_style, note=HYPOTHETICAL_N_NOTE)
        story.append(Spacer(1, 8))

        story.extend(
            [
                Paragraph("Teste estatístico do cenário hipotético", subheading_style),
                Paragraph(scenario.get("statistical_text") or "", body_style),
                Spacer(1, 8),
            ]
        )
        _append_pdf_table_intro(story, "hypothetical_tests", body_style, caption_style)
        test_rows = [["Radionuclídeo", "Compart.", "n", "Shapiro", "Teste", "p-value", "P95/RL", "Acima", "Conclusão"]]
        for result in scenario.get("statistical_tests", []):
            test_rows.append(
                [
                    result.get("radionuclide", ""),
                    result.get("compartment", ""),
                    str(result.get("n") or "—"),
                    result.get("shapiro_p_text") or "—",
                    result.get("test_label") or "—",
                    result.get("p_value_text") or "—",
                    result.get("p95_ratio_text") or "—",
                    str(result.get("exceedance_count") or 0),
                    result.get("conclusion") or "",
                ]
            )
        test_table = Table(
            test_rows,
            colWidths=[1.5 * cm, 1.7 * cm, 0.8 * cm, 1.2 * cm, 2.0 * cm, 1.1 * cm, 1.2 * cm, 1.0 * cm, 4.0 * cm],
            repeatRows=1,
        )
        test_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eaf1f3")),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#c9d7dc")),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 7),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        story.append(test_table)
        _append_pdf_table_notes(story, "hypothetical_tests", note_style, note=HYPOTHETICAL_TEST_NOTE)
        story.append(Spacer(1, 8))

    empirical = scenario.get("empirical_activity_statistics") or {}
    if empirical:
        story.extend(
            [
                Paragraph("Dados reais de atividade total TAR", subheading_style),
                Paragraph(_pdf_text(empirical.get("narrative_text") or ""), body_style),
                Paragraph(_pdf_text(" | ".join(f"{card.get('label')}: {card.get('value')}" for card in empirical.get("cards") or [])), body_style),
                Spacer(1, 8),
            ]
        )
        _append_pdf_table_intro(story, "empirical_groups", body_style, caption_style)
        group_rows = [["Grupo", "Amostras", "A-TAR det.", "A-TAR < MDA", "Ausente", "Média", "Mediana", "P95"]]
        for item in empirical.get("groups") or []:
            group_rows.append(
                [
                    item.get("group", ""),
                    str(item.get("sample_count") or 0),
                    str(item.get("activity_detected_count") or 0),
                    str(item.get("activity_censored_count") or 0),
                    str(item.get("activity_missing_count") or 0),
                    item.get("activity_mean_text", "—"),
                    item.get("activity_median_text", "—"),
                    item.get("activity_p95_text", "—"),
                ]
            )
        group_table = Table(group_rows, colWidths=[2.7 * cm, 1.4 * cm, 1.5 * cm, 1.7 * cm, 1.4 * cm, 2.2 * cm, 2.2 * cm, 2.2 * cm], repeatRows=1)
        group_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eaf1f3")),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#c9d7dc")),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 6.7),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        story.append(group_table)
        _append_pdf_table_notes(story, "empirical_groups", note_style, note=empirical.get("mda_policy") or EMPIRICAL_ACTIVITY_NOTE)
        story.append(Spacer(1, 8))
        _append_pdf_table_intro(story, "empirical_radionuclides", body_style, caption_style)

        radionuclide_rows = [["Grupo", "Radionuclídeo", "Status", "Amostras", "Detect.", "< MDA", "Aus.", "Taxa", "Média", "Mediana", "P95"]]
        for item in empirical.get("radionuclide_rows") or []:
            radionuclide_rows.append(
                [
                    item.get("group", ""),
                    item.get("radionuclide", ""),
                    item.get("model_status", ""),
                    str(item.get("sample_count") or 0),
                    str(item.get("detected_count") or 0),
                    str(item.get("censored_count") or 0),
                    str(item.get("missing_count") or 0),
                    item.get("detected_rate_text", "—"),
                    item.get("mean_text", "—"),
                    item.get("median_text", "—"),
                    item.get("p95_text", "—"),
                ]
            )
        radionuclide_table = Table(
            radionuclide_rows,
            colWidths=[2.0 * cm, 1.2 * cm, 1.9 * cm, 1.0 * cm, 1.0 * cm, 1.0 * cm, 0.9 * cm, 1.1 * cm, 1.8 * cm, 1.8 * cm, 1.8 * cm],
            repeatRows=1,
        )
        radionuclide_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eaf1f3")),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#c9d7dc")),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 5.8),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        story.append(radionuclide_table)
        _append_pdf_table_notes(
            story,
            "empirical_radionuclides",
            note_style,
            note="Nb-95 permanece como observado não modelado porque não há linha equivalente na fórmula da planilha TAR atual.",
        )
        story.append(Spacer(1, 8))
        _append_pdf_table_intro(story, "empirical_modeled", body_style, caption_style)

        modeled_rows = [["Grupo", "Radionuclídeo", "Compart.", "n", "Média", "Mediana", "P95", "RL", "P95/RL", "Status", "Freq. > RL"]]
        for item in empirical.get("modeled_compartment_rows") or []:
            modeled_rows.append(
                [
                    item.get("group", ""),
                    item.get("radionuclide", ""),
                    item.get("compartment", ""),
                    str(item.get("n") or "—"),
                    item.get("mean_text", "—"),
                    item.get("median_text", "—"),
                    item.get("p95_text", "—"),
                    item.get("report_level_text", "—"),
                    item.get("report_level_p95_ratio_text", "—"),
                    item.get("report_level_status", "—"),
                    item.get("report_level_exceedance_rate_text", "—"),
                ]
            )
        modeled_table = Table(
            modeled_rows,
            colWidths=[1.7 * cm, 1.1 * cm, 1.4 * cm, 0.7 * cm, 1.5 * cm, 1.5 * cm, 1.5 * cm, 1.5 * cm, 1.1 * cm, 1.1 * cm, 1.5 * cm],
            repeatRows=1,
        )
        modeled_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eaf1f3")),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#c9d7dc")),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 5.2),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        story.append(modeled_table)
        _append_pdf_table_notes(
            story,
            "empirical_modeled",
            note_style,
            note="As frações Si foram calculadas por amostra a partir dos radionuclídeos detectados; < MDA> não participa do denominador.",
        )
        story.append(Spacer(1, 8))

        _append_pdf_table_intro(story, "empirical_inferential", body_style, caption_style)
        inferential_rows = [["Grupo", "Radionuclídeo", "Compart.", "n", "RL", "P95 razão", "Ultrap.", "Freq.", "IC95% freq.", "Teste", "p-value", "Conclusão"]]
        for item in empirical.get("inferential_rows") or []:
            inferential_rows.append(
                [
                    item.get("group", ""),
                    item.get("radionuclide", ""),
                    item.get("compartment", ""),
                    str(item.get("n") or "—"),
                    item.get("reference_value_text", "—"),
                    item.get("p95_ratio_text", "—"),
                    str(item.get("exceedance_count") if item.get("exceedance_count") is not None else "—"),
                    item.get("exceedance_rate_text", "—"),
                    item.get("exceedance_ci95_text", "—"),
                    item.get("test_label", "—"),
                    item.get("p_value_text", "—"),
                    item.get("conclusion") or item.get("reason", ""),
                ]
            )
        inferential_table = apply_table_style(
            Table(
                inferential_rows,
                colWidths=[1.4 * cm, 1.1 * cm, 1.3 * cm, 0.6 * cm, 1.3 * cm, 1.0 * cm, 0.9 * cm, 0.9 * cm, 1.5 * cm, 1.7 * cm, 1.0 * cm, 2.5 * cm],
                repeatRows=1,
            ),
            font_size=3.8,
        )
        story.append(inferential_table)
        _append_pdf_table_notes(
            story,
            "empirical_inferential",
            note_style,
            note="O Report Level é referência fixa; a incerteza estatística é estimada a partir das amostras reais calculadas pela fórmula.",
        )
        story.append(Spacer(1, 8))

    statistical = scenario.get("statistical_comparison") or {}
    if statistical:
        story.extend(
            [
                Paragraph("Estatística calculado vs ERICA vs norma", subheading_style),
                Paragraph(_pdf_text(statistical.get("narrative_text") or ""), body_style),
                Paragraph(
                    "Esta seção não usa replicações aleatórias. A norma é referência fixa; os valores do ERICA Tool são estimativas para visualização até substituição por saídas reais.",
                    note_style,
                ),
                Spacer(1, 8),
            ]
        )

        for rows_key, table_key in [("calculated_rows", "stat_calculated"), ("erica_rows", "stat_erica"), ("norm_rows", "stat_norms")]:
            _append_pdf_table_intro(story, table_key, body_style, caption_style)
            if rows_key == "norm_rows":
                stat_source_rows = [["Radionuclídeo", "Compart.", "Referência", "Valor normativo", "Unidade"]]
                for item in statistical.get(rows_key) or []:
                    stat_source_rows.append(
                        [
                            item.get("radionuclide", ""),
                            item.get("compartment", ""),
                            item.get("reference", ""),
                            item.get("value_text", "—"),
                            item.get("unit", ""),
                        ]
                    )
            else:
                stat_source_rows = [["Radionuclídeo", "Compart.", "Valor base", "Unidade", "Origem"]]
                for item in statistical.get(rows_key) or []:
                    stat_source_rows.append(
                        [
                            item.get("radionuclide", ""),
                            item.get("compartment", ""),
                            item.get("value_text", "—"),
                            item.get("unit", ""),
                            item.get("method", ""),
                        ]
                    )
            stat_source_table = apply_table_style(
                Table(stat_source_rows, colWidths=[2.0 * cm, 2.1 * cm, 2.6 * cm, 2.1 * cm, 7.0 * cm], repeatRows=1),
                font_size=5.5,
            )
            story.append(stat_source_table)
            _append_pdf_table_notes(
                story,
                table_key,
                note_style,
                note="Os valores normativos são fixos; os valores calculados e ERICA são bases determinísticas ou estimadas, sem replicações aleatórias.",
            )
            story.append(Spacer(1, 8))

        _append_pdf_table_intro(story, "stat_descriptive", body_style, caption_style)
        descriptive_rows = [["Conjunto", "Radionuclídeo", "Compart.", "n", "Média", "Mediana", "DP", "Q1", "Q3", "P95", "CV"]]
        for item in statistical.get("descriptive_rows") or []:
            descriptive_rows.append(
                [
                    item.get("dataset", ""),
                    item.get("radionuclide", ""),
                    item.get("compartment", ""),
                    str(item.get("n") or "—"),
                    item.get("mean_text", "—"),
                    item.get("median_text", "—"),
                    item.get("stdev_text", "—"),
                    item.get("q1_text", "—"),
                    item.get("q3_text", "—"),
                    item.get("p95_text", "—"),
                    item.get("cv_text", "—"),
                ]
            )
        descriptive_table = apply_table_style(
            Table(
                descriptive_rows,
                colWidths=[1.5 * cm, 1.4 * cm, 1.5 * cm, 0.8 * cm, 1.4 * cm, 1.4 * cm, 1.3 * cm, 1.1 * cm, 1.1 * cm, 1.3 * cm, 1.1 * cm],
                repeatRows=1,
            ),
            font_size=4.6,
        )
        story.append(descriptive_table)
        _append_pdf_table_notes(story, "stat_descriptive", note_style, note="A estatística usa os valores disponíveis por radionuclídeo e compartimento; não há aumento artificial de N por sorteio.")
        story.append(Spacer(1, 8))

        _append_pdf_table_intro(story, "stat_paired", body_style, caption_style)
        paired_rows = [["Escopo", "Compart.", "n", "Mediana calc./ERICA", "IC95% razão", "Teste", "p-value", "Conclusão"]]
        for item in statistical.get("paired_comparison_rows") or []:
            paired_rows.append(
                [
                    item.get("scope") or item.get("radionuclide", ""),
                    item.get("compartment", ""),
                    str(item.get("n") or "—"),
                    item.get("median_ratio_text", "—"),
                    f"{item.get('ci95_low_text', '—')} - {item.get('ci95_high_text', '—')}",
                    item.get("test_label", "—"),
                    item.get("p_value_text", "—"),
                    item.get("conclusion") or item.get("reason", ""),
                ]
            )
        paired_table = apply_table_style(
            Table(paired_rows, colWidths=[2.0 * cm, 1.5 * cm, 0.7 * cm, 1.7 * cm, 2.0 * cm, 2.0 * cm, 1.1 * cm, 4.0 * cm], repeatRows=1),
            font_size=4.5,
        )
        story.append(paired_table)
        _append_pdf_table_notes(story, "stat_paired", note_style, note="A comparação pareada usa log(calculado / ERICA) e tem finalidade exploratória; os valores do ERICA são estimativas visuais.")
        story.append(Spacer(1, 8))

    sensitivity = scenario.get("sensitivity") or {}
    if sensitivity:
        story.extend(
            [
                Paragraph("Análise de sensibilidade Monte Carlo", subheading_style),
                Paragraph(sensitivity.get("narrative_text") or "", body_style),
                Paragraph(" | ".join(f"{card.get('label')}: {card.get('value')}" for card in sensitivity.get("cards") or []), body_style),
                Spacer(1, 8),
            ]
        )
        _append_pdf_table_intro(story, "sensitivity_variables", body_style, caption_style)
        variable_rows = [["Variável", "Distribuição", "Parâmetros", "Base", "Unidade", "Uso no modelo"]]
        for variable in sensitivity.get("variables") or []:
            variable_rows.append(
                [
                    variable.get("label", ""),
                    variable.get("distribution", ""),
                    variable.get("parameters", ""),
                    variable.get("base_value_text", "—"),
                    variable.get("unit", ""),
                    variable.get("description", ""),
                ]
            )
        variable_table = Table(variable_rows, colWidths=[3.1 * cm, 2.1 * cm, 3.2 * cm, 2.1 * cm, 1.7 * cm, 4.1 * cm], repeatRows=1)
        variable_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eaf1f3")),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#c9d7dc")),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 6.5),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        story.append(variable_table)
        _append_pdf_table_notes(story, "sensitivity_variables", note_style, note=sensitivity.get("source_note") or SENSITIVITY_NOTE)
        story.append(Spacer(1, 8))

        for title, image_buffer, width, height in _sensitivity_chart_images(sensitivity):
            image_width = 16.4 * cm
            story.append(Paragraph(CHART_EXPLANATIONS.get(title, ""), body_style))
            story.append(Paragraph(title, caption_style))
            story.append(RLImage(image_buffer, width=image_width, height=image_width * (height / width)))
            story.append(Spacer(1, 8))

        influence_rows = [["Variável", "Spearman", "|Correlação|", "Sentido"]]
        for item in sensitivity.get("influence_rows") or []:
            influence_rows.append(
                [
                    item.get("label", ""),
                    item.get("correlation_text", "—"),
                    item.get("absolute_correlation_text", "—"),
                    item.get("direction", "—"),
                ]
            )
        influence_table = Table(influence_rows, colWidths=[7.0 * cm, 3.0 * cm, 3.0 * cm, 3.0 * cm], repeatRows=1)
        influence_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eaf1f3")),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#c9d7dc")),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 7),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        _append_pdf_table_intro(story, "sensitivity_influence", body_style, caption_style)
        story.append(influence_table)
        _append_pdf_table_notes(
            story,
            "sensitivity_influence",
            note_style,
            note=(
                "Correlação positiva indica que valores maiores da variável tendem a aumentar a maior razão contra o Report Level; "
                "correlação negativa indica que valores maiores tendem a reduzir essa razão."
            ),
        )
        story.append(Spacer(1, 8))

        result_rows = [["Radionuclídeo", "Compart.", "Média", "Mín.", "Máx.", "P95", "RL", "P95/RL", "Prob. > RL"]]
        for item in sensitivity.get("summary_rows") or []:
            result_rows.append(
                [
                    item.get("radionuclide", ""),
                    item.get("compartment", ""),
                    item.get("mean_text", "—"),
                    item.get("min_text", "—"),
                    item.get("max_text", "—"),
                    item.get("p95_text", "—"),
                    item.get("report_level_text", "—"),
                    item.get("p95_report_level_ratio_text", "—"),
                    item.get("exceedance_probability_text", "—"),
                ]
            )
        result_table = Table(
            result_rows,
            colWidths=[1.4 * cm, 1.6 * cm, 1.8 * cm, 1.8 * cm, 1.8 * cm, 1.8 * cm, 1.8 * cm, 1.2 * cm, 1.3 * cm],
            repeatRows=1,
        )
        result_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eaf1f3")),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#c9d7dc")),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 5.8),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        _append_pdf_table_intro(story, "sensitivity_results", body_style, caption_style)
        story.append(result_table)
        _append_pdf_table_notes(
            story,
            "sensitivity_results",
            note_style,
            note=(
                "A probabilidade empírica é calculada como número de simulações acima do Report Level dividido pelo número total de simulações. "
                "Quando não há Report Level cadastrado, a probabilidade fica sem referência."
            ),
        )
        story.append(Spacer(1, 8))

    story.extend(
        [
            Paragraph("Suficiência estatística", subheading_style),
            Paragraph(summary["inferential_assessment"]["reason"], body_style),
            Spacer(1, 8),
        ]
    )
    _append_pdf_table_intro(story, "minimums", body_style, caption_style)

    minimum_rows = [["Teste", "Mínimo técnico", "Recomendado para relatório"]]
    for item in INFERENTIAL_TEST_MINIMUMS:
        minimum_rows.append([item["test"], item["technical_minimum"], item["recommended"]])
    table = Table(minimum_rows, colWidths=[4.0 * cm, 5.0 * cm, 7.0 * cm], repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eaf1f3")),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#c9d7dc")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    story.append(table)
    _append_pdf_table_notes(story, "minimums", note_style, note=DETERMINISTIC_N_NOTE)
    document.build(story)
    filename = f"relatorio_tar_{summary['selected_scenario']}.pdf"
    return output.getvalue(), filename, PDF_MIME
